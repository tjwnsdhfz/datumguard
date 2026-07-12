from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = ROOT / "docs" / "awards-2026" / "COMPETITION_DESCRIPTION.md"
DEFAULT_OUTPUT = ROOT / "docs" / "awards-2026" / "BIM_AWARDS_2026_STUDENT_RESEARCH_DESCRIPTION.docx"

ASCII_FONT = "Calibri"
EAST_ASIA_FONT = "Malgun Gothic"
INK = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(84, 96, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "C9D4E2"

# Official-form overrides to the narrative_proposal preset.
PAGE_WIDTH_DXA = 11907  # A4
MARGIN_DXA = 794  # 14 mm
CONTENT_WIDTH_DXA = PAGE_WIDTH_DXA - (2 * MARGIN_DXA)
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 60, "bottom": 60, "start": 120, "end": 120}


def set_run_font(
    run,
    *,
    size: float,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    ascii_font: str = ASCII_FONT,
    east_asia_font: str = EAST_ASIA_FONT,
) -> None:
    run.font.name = ascii_font
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)


def set_style_font(style, *, size: float, color: RGBColor, bold: bool) -> None:
    style.font.name = ASCII_FONT
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), ASCII_FONT)
    r_fonts.set(qn("w:hAnsi"), ASCII_FONT)
    r_fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, size=10, color=INK, bold=False)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.6
    normal.paragraph_format.widow_control = True

    title = doc.styles["Title"]
    set_style_font(title, size=17, color=INK, bold=True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.line_spacing = 1.05
    title.paragraph_format.keep_with_next = True
    title_p_pr = title.element.get_or_add_pPr()
    title_border = title_p_pr.find(qn("w:pBdr"))
    if title_border is not None:
        title_p_pr.remove(title_border)

    subtitle = doc.styles["Subtitle"]
    set_style_font(subtitle, size=9.5, color=MUTED, bold=False)
    subtitle.font.italic = False
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(6)
    subtitle.paragraph_format.line_spacing = 1.15
    subtitle.paragraph_format.keep_with_next = True

    heading_1 = doc.styles["Heading 1"]
    set_style_font(heading_1, size=13.5, color=INK, bold=True)
    heading_1.paragraph_format.space_before = Pt(0)
    heading_1.paragraph_format.space_after = Pt(5)
    heading_1.paragraph_format.line_spacing = 1.1
    heading_1.paragraph_format.keep_with_next = True

    heading_2 = doc.styles["Heading 2"]
    set_style_font(heading_2, size=11.2, color=BLUE, bold=True)
    heading_2.paragraph_format.space_before = Pt(5)
    heading_2.paragraph_format.space_after = Pt(2)
    heading_2.paragraph_format.line_spacing = 1.1
    heading_2.paragraph_format.keep_with_next = True


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in CELL_MARGIN_DXA.items():
        tag = "w:start" if side == "start" else "w:end" if side == "end" else f"w:{side}"
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    total = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Cm(width / 1440 * 2.54)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), BORDER)


def add_inline_markdown(paragraph, text: str, *, size: float = 10) -> None:
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, color=INK, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(
                run,
                size=size,
                color=INK,
                ascii_font="Consolas",
                east_asia_font=EAST_ASIA_FONT,
            )
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=INK)


def add_callout(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.35)
    paragraph.paragraph_format.right_indent = Cm(0.2)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.6
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), "2E74B5")
    left.set(qn("w:space"), "8")
    borders.append(left)
    p_pr.append(borders)
    add_inline_markdown(paragraph, text, size=10)


def normalize_paragraph(lines: list[str]) -> str:
    return " ".join(line.strip() for line in lines).strip()


def render_section_body(doc: Document, body: str) -> None:
    paragraph_lines: list[str] = []
    quote_lines: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_lines:
            return
        text = normalize_paragraph(paragraph_lines)
        paragraph_lines.clear()
        if not text:
            return
        if text.startswith("`") and text.endswith("`"):
            add_callout(doc, text)
            return
        paragraph = doc.add_paragraph()
        add_inline_markdown(paragraph, text)

    def flush_quote() -> None:
        if not quote_lines:
            return
        add_callout(doc, normalize_paragraph(quote_lines))
        quote_lines.clear()

    for raw_line in body.splitlines():
        line = raw_line.rstrip()
        if line.startswith("### "):
            flush_paragraph()
            flush_quote()
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
        elif line.startswith(">"):
            flush_paragraph()
            quote_lines.append(line[1:].strip())
        elif not line.strip():
            flush_paragraph()
            flush_quote()
        elif line.startswith("- "):
            flush_paragraph()
            flush_quote()
            raise ValueError("Bulleted source content requires a real numbering definition")
        else:
            flush_quote()
            paragraph_lines.append(line)

    flush_paragraph()
    flush_quote()


def section_text(source: str, page: int) -> tuple[str, str]:
    pattern = re.compile(
        rf"^## {page}쪽\. (?P<title>.+?)\r?\n(?P<body>.*?)(?=^---\r?$|\Z)",
        flags=re.MULTILINE | re.DOTALL,
    )
    match = pattern.search(source)
    if match is None:
        raise ValueError(f"Could not find page {page} section")
    return match.group("title").strip(), match.group("body").strip()


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction_node = OxmlElement("w:instrText")
    instruction_node.set(qn("xml:space"), "preserve")
    instruction_node.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction_node, separate, text, end])
    set_run_font(run, size=8.5, color=MUTED)


def configure_page(doc: Document) -> None:
    doc.settings.odd_and_even_pages_header_footer = True
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    for header in (
        section.first_page_header,
        section.even_page_header,
        section.header,
    ):
        header_p = header.paragraphs[0]
        header_p.text = ""
        header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_p.paragraph_format.space_after = Pt(0)
        run = header_p.add_run("DATUMGUARD  /  OPENBIM EVIDENCE GUARD")
        set_run_font(run, size=8, color=MUTED, bold=True)

    for footer in (
        section.first_page_footer,
        section.even_page_footer,
        section.footer,
    ):
        footer_p = footer.paragraphs[0]
        footer_p.text = ""
        footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_p.paragraph_format.space_before = Pt(0)
        run = footer_p.add_run("BIM AWARDS 2026 · STUDENT RESEARCH  |  ")
        set_run_font(run, size=8.5, color=MUTED)
        add_field(footer_p, "PAGE")
        run = footer_p.add_run(" / 3")
        set_run_font(run, size=8.5, color=MUTED)


def add_first_page_masthead(doc: Document) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(2)
    run = kicker.add_run("BIM AWARDS 2026  ·  학생부 RESEARCH")
    set_run_font(run, size=9, color=BLUE, bold=True)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("OpenBIM Evidence Guard")
    set_run_font(run, size=17, color=INK, bold=True)

    subtitle = doc.add_paragraph(style="Subtitle")
    add_inline_markdown(
        subtitle,
        "IDS 기반 가상 FAB Utility IFC의 정보요구조건 및 리비전 무결성 독립 재검증 연구",
        size=10.2,
    )

    english = doc.add_paragraph(style="Subtitle")
    add_inline_markdown(
        english,
        "Independent Revalidation of IDS-Based Information Requirements and "
        "Revision Integrity for Virtual FAB Utility IFC Models",
        size=8.8,
    )

    rows = [
        ("부문·분야", "학생 / Research"),
        ("소속·참가자", "[학교명] [학과명]  ·  [학년] [성명]"),
        ("주요 BIM 도구", "IFC4 · IfcOpenShell 0.8.5 · IfcTester 0.8.5 · JSON/HTML evidence"),
        ("판정 경계", "Synthetic research validation only · 설계·안전·법규·제작·시공 승인 아님"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [1850, CONTENT_WIDTH_DXA - 1850])
    set_table_borders(table)
    for row_index, (label, value) in enumerate(rows):
        label_cell, value_cell = table.rows[row_index].cells
        set_cell_shading(label_cell, LIGHT_BLUE)
        set_cell_shading(value_cell, "FFFFFF")
        label_p = label_cell.paragraphs[0]
        label_p.paragraph_format.space_before = Pt(0)
        label_p.paragraph_format.space_after = Pt(0)
        label_p.paragraph_format.line_spacing = 1.1
        run = label_p.add_run(label)
        set_run_font(run, size=10, color=INK, bold=True)
        value_p = value_cell.paragraphs[0]
        value_p.paragraph_format.space_before = Pt(0)
        value_p.paragraph_format.space_after = Pt(0)
        value_p.paragraph_format.line_spacing = 1.1
        run = value_p.add_run(value)
        set_run_font(run, size=10, color=INK)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def build_document(source_path: Path, output_path: Path) -> None:
    source = source_path.read_text(encoding="utf-8")
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    add_first_page_masthead(doc)

    for page in (1, 2, 3):
        title, body = section_text(source, page)
        heading = doc.add_paragraph(f"{page}. {title}", style="Heading 1")
        if page > 1:
            heading.paragraph_format.page_break_before = True
        render_section_body(doc, body)

    properties = doc.core_properties
    properties.title = "BIM Awards 2026 Student Research Description"
    properties.subject = "OpenBIM Evidence Guard competition description"
    properties.author = "DatumGuard Student Research"
    properties.keywords = "BIM Awards, OpenBIM, IFC, IDS, student research"
    properties.comments = "Generated from docs/awards-2026/COMPETITION_DESCRIPTION.md"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build the BIM Awards 2026 description DOCX.")
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    build_document(args.source.resolve(), args.output.resolve())
    print(args.output.resolve())


if __name__ == "__main__":
    main()
